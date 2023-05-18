import random
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENTATION
import os
from tqdm import tqdm
import sys


def generate_random_number():
    return random.randint(1, 20)


def generate_random_operator():
    return random.choice(['+', '-'])


def generate_addition_question(num):
    operand1 = generate_random_number()
    operand2 = generate_random_number()
    return f"{num:02d}. {operand1} + {operand2} ="


def generate_subtraction_question(num):
    operand1 = generate_random_number()
    operand2 = generate_random_number()
    return f"{num:02d}. {operand1} - {operand2} ="


def generate_multiplication_question(num):
    operand1 = random.randint(0, 12)
    operand2 = random.randint(0, 12)
    return f"{num:02d}. {operand1} x {operand2} ="


def generate_division_question(num):
    divisor = random.randint(3, 12)
    dividend = divisor * random.randint(1, 12)
    return f"{num:02d}. {dividend} ÷ {divisor} ="


def generate_questions():
    questions = []
    num = 1

    for _ in range(4):
        for _ in range(5):
            operator = generate_random_operator()
            if operator == '+':
                questions.append(generate_addition_question(num))
            else:
                questions.append(generate_subtraction_question(num))
            num += 1

        for _ in range(10):
            equation_parts = []
            for _ in range(3):
                equation_parts.append(str(generate_random_number()))
                equation_parts.append(generate_random_operator())
            equation_parts.pop()  # Remove the last operator
            equation = ' '.join(equation_parts)
            questions.append(f"{num:02d}. {equation} =")
            num += 1

        for _ in range(6):
            questions.append(generate_multiplication_question(num))
            num += 1

        for _ in range(4):
            questions.append(generate_division_question(num))
            num += 1

    return questions


def generate_word_document(times):
    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Cordia New'
    font.size = Pt(13.5)

    section = document.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(0)

    column_count = 4

    if times < 20:
        for time in range(times):
            questions = generate_questions()

            table = document.add_table(rows=25, cols=column_count)
            table.autofit = False
            table.allow_autofit = False

            for i, question in enumerate(tqdm(questions, desc=f"Generating Questions: {time+1}/{times}", ncols=80)):
                row = i % 25
                column = i // 25
                cell = table.cell(row, column)
                cell.text = question
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)

            if time != times - 1:
                document.add_page_break()
    else:
        for time in tqdm (range (times), desc="Generating Questions", ncols=80):
            questions = generate_questions()

            table = document.add_table(rows=25, cols=column_count)
            table.autofit = False
            table.allow_autofit = False

            for i, question in enumerate(questions):
                row = i % 25
                column = i // 25
                cell = table.cell(row, column)
                cell.text = question
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)

            if time != times - 1:
                document.add_page_break()

    directory = "Generated_Problems"
    if not os.path.exists(directory):
        try:
            os.makedirs(directory)
        except Exception as e:
            print("\033[91m\nA critical error occurred. Please contact the programmer.")
            print("Error:", str(e), "\033[0m")
            sys.exit(1)

    filename = generate_file_name(times, directory)
    document.save(filename)
    print(f"\033[92m\nDocument '{filename[19:]}' saved successfully at '{filename}'.\033[0m")


def generate_file_name(times, directory):
    base_filename = f"{times*100}_questions"
    filename = f"{directory}/{base_filename}_1.docx"

    count = 1
    while os.path.exists(filename):
        count += 1
        filename = f"{directory}/{base_filename}_{count}.docx"
    filename = f"{base_filename}_{count}.docx"
    
    if count == 2:
        print(f"\033[93m\nFilename '{base_filename}_1.docx' already exists. Changing the filename to '{filename}'.\033[0m")
    elif count == 3:
        print(f"""\033[93m\nFilename '{base_filename}_1.docx' and '{base_filename}_2.docx' already exist.
    Changing the filename to '{base_filename}_3.docx'.\033[0m""")
    elif count > 3:
        print(f"""\033[93m\nFilename '{base_filename}_1.docx' to '{base_filename}_{count-1}.docx' already exist.
    Changing the filename to '{filename}'.\033[0m""")
    
    filename = f"{directory}/{base_filename}_{count}.docx"

    return filename


def main():
    print("\033[92m\nThis application will generate as many simple math questions as you want.")
    print("There will be one hundred questions on one page.\033[93m")

    directory = "Generated_Problems"
    if os.path.exists(directory):
        print("\033[93m\nFolder 'Generated_Problems' already exists. The result will be inside the folder.\033[0m")

    times = int(input("\033[96m\nPlease input the desired number of pages> \033[0m"))

    generate_word_document(times)

    input("\033[95m\nPress Enter to exit...\033[0m")


if __name__ == '__main__':
    main()
