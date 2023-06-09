import os
import file_name as nameGen
from tqdm import tqdm
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENTATION


def generate_word_document(times, directory, all_ques):
    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Cordia New'
    font.size = Pt(13.5)

    section = document.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(0)

    column_count = 4

    if times < 20:
        for time in range(times):
            questions = all_ques[time]

            table = document.add_table(rows=25, cols=column_count)
            table.autofit = False
            table.allow_autofit = False

            for i, question in enumerate(tqdm(questions, desc=f"Generating Questions: {time+1}/{times}", ncols=80)):
                row = i % 25
                column = i // 25
                cell = table.cell(row, column)
                cell.text = question
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)

            if time != times - 1:
                document.add_page_break()
    else:
        for time in tqdm (range (times), desc="Generating Questions", ncols=80):
            questions = all_ques[time]

            table = document.add_table(rows=25, cols=column_count)
            table.autofit = False
            table.allow_autofit = False

            for i, question in enumerate(questions):
                row = i % 25
                column = i // 25
                cell = table.cell(row, column)
                cell.text = question
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)

            if time != times - 1:
                document.add_page_break()


    filename = nameGen.generate_file_name(times, directory)
    document.save(filename)
    print(f"\033[92m\nDocument '{filename[19:]}' saved successfully at '{filename}'.\033[0m")